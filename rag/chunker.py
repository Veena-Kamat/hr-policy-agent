"""
rag/chunker.py
Document text extraction and chunking pipeline.
Supports: PDF (pdfplumber), DOCX (python-docx), TXT
"""

import re

CHUNK_SIZE    = 500   # target words per chunk
CHUNK_OVERLAP = 75    # words of overlap between adjacent chunks


def extract_and_chunk(filepath: str, ext: str) -> list[dict]:
    """
    Extract text from a file and return a list of chunk dicts:
        { 'content': str, 'page': int|None, 'index': int }
    """
    ext = ext.lower()

    if ext == '.pdf':
        text_pages = _extract_pdf(filepath)
    elif ext == '.docx':
        text_pages = _extract_docx(filepath)
    elif ext == '.txt':
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        text_pages = [{'page': None, 'text': text}]
    else:
        raise ValueError(f'Unsupported extension: {ext}')

    return _chunk_pages(text_pages)


# ── Extractors ───────────────────────────────────────────────────────────────

def _extract_pdf(filepath: str) -> list[dict]:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required for PDF support: pip install pdfplumber")

    pages = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ''
            if text.strip():
                pages.append({'page': i, 'text': text})
    return pages


def _extract_docx(filepath: str) -> list[dict]:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX support: pip install python-docx")

    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = '\n'.join(paragraphs)
    return [{'page': None, 'text': full_text}]


# ── Chunker ──────────────────────────────────────────────────────────────────

def _chunk_pages(text_pages: list[dict]) -> list[dict]:
    chunks = []
    chunk_index = 0

    for page_info in text_pages:
        text = _clean_text(page_info['text'])
        words = text.split()

        if not words:
            continue

        i = 0
        while i < len(words):
            chunk_words = words[i: i + CHUNK_SIZE]
            content = ' '.join(chunk_words).strip()
            if len(content) > 50:   # skip trivially small chunks
                chunks.append({
                    'content': content,
                    'page': page_info.get('page'),
                    'index': chunk_index
                })
                chunk_index += 1
            i += CHUNK_SIZE - CHUNK_OVERLAP

    return chunks


def _clean_text(text: str) -> str:
    """Normalise whitespace and remove non-printable characters."""
    text = re.sub(r'[^\x20-\x7E\u00A0-\uFFFF\n]', ' ', text)
    text = re.sub(r'\n+', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()
